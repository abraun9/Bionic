import os
from docx import Document

def bold_half_of_word(word):
    half_index = len(word) // 2
    return (word[:half_index], word[half_index:])

def bionic_write(input_filename, output_filename):
    # Load the original document
    doc = Document(input_filename)
    bionic_doc = Document()

    for para in doc.paragraphs:
        bionic_para = bionic_doc.add_paragraph()
        for word in para.text.split():
            bold_part, normal_part = bold_half_of_word(word)
            run = bionic_para.add_run(bold_part)
            run.bold = True
            bionic_para.add_run(normal_part + ' ')

    bionic_doc.save(output_filename)

def process_file(input_filename):
    # Check if the file has the correct extension
    if not input_filename.endswith('.docx'):
        raise ValueError('File must have a .docx extension.')
    # Check if the file exists
    if not os.path.isfile(input_filename):
        raise FileNotFoundError(f'The file {input_filename} does not exist in the directory.')

    # Generate the output file name
    output_filename = input_filename.replace('.docx', '_bionic.docx')
    # Perform the bionic writing conversion
    bionic_write(input_filename, output_filename)

def main():
    try:
        # Prompt the user for the input file name
        input_filename = input('Enter the name of the Word file (with .docx): ')
        # Process the file
        process_file(input_filename)
        print(f'Bionic writing conversion complete. Saved as {input_filename.replace(".docx", "_bionic.docx")}')
    except ValueError as ve:
        print(f'Error: {ve}')
    except FileNotFoundError as fnfe:
        print(f'Error: {fnfe}')
    except Exception as e:
        print(f'An unexpected error occurred: {e}')

if __name__ == '__main__':
    main()